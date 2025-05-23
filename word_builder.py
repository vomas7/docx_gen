import decimal
from collections import namedtuple
from copy import deepcopy
from decimal import Decimal
from typing import Any

from docx import Document, oxml
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm
from docx.table import Table, _Cell, _Row
from docx.text.paragraph import Paragraph
from lxml import etree

from doc import DOC
from helpers.cell_selector import CellSelection
from helpers.styler import styler


class WordBuilder(DOC):
    """Класс для работы создания и работы с объектами документа."""

    data: list
    _data: list
    table: Table

    def __init__(self):
        super().__init__()

    def add_table(
        self,
        table_data: list[list],
        rows: int = None,
        cols: int = None,
        table_style: str = None,
    ) -> Table:
        """
        Метод добавляет таблицу в документ и перезаписывает переменную table.
        После добавления новой таблицы, все последующие методы будут работать
        только с ней. Кроме создания происходит и заполнение таблицы данными.
        """
        self.data = table_data
        self._data = sum(self.data, []) if self.data else []
        self.table = self._build_table(rows=rows, cols=cols, table_style=table_style)
        self.table.autofit = False
        self.table.allow_autofit = False
        self._fill_table()
        return self.table

    @styler.style_text()
    def add_text(self, text: str, **kwargs) -> tuple[Paragraph, dict]:
        """
        Метод добавляет переданную строку как текст документа,
        после чего styler применяет к нему переданные стили.
        """
        if self.doc.paragraphs[0].runs[0].text == "123":  ############
            self.doc.paragraphs[0].runs[0].text = text
            paragraph = self.doc.paragraphs[0]
        else:
            paragraph = self.doc.add_paragraph(text)
        return paragraph, kwargs

    def insert_picture(
        self, path: str, height: int = 2.06, width: int = 13.95, align: str = "center"
    ) -> None:
        if self.doc.paragraphs[0].runs[0].text == "123":  ###########
            self.doc.paragraphs[0].runs[0].text = ""
            paragraph = self.doc.paragraphs[0]
            run = paragraph.runs[0]
        else:
            paragraph = self.doc.add_paragraph()
            run = paragraph.add_run()
        shape = run.add_picture(path)
        shape.width = Cm(width)
        shape.height = Cm(height)
        match align:
            case "left":
                paragraph.alignment = 0
            case "right":
                paragraph.alignment = 2
            case "center":
                paragraph.alignment = 1

    @property
    def last_row(self) -> int:
        """Возвращает индекс последней строки."""
        return len(self.table.rows) - 1

    @property
    def last_col(self) -> int:
        """Возвращает индекс последней колонки."""
        return len(self.table.columns) - 1

    @staticmethod
    def iter_cell_text(cell: _Cell) -> iter:
        """Итератор по тексту ячейки."""
        yield from cell.paragraphs

    def get_row(self, row_index: int) -> _Row:
        """Геттер, возвращает объект строки по индексу."""
        return self.table.rows[row_index]

    @property
    def rows(self) -> iter:
        """Итератор по строкам таблицы."""
        yield from self.table.rows

    @property
    def cells(self) -> iter:
        """Итератор по всем ячейкам."""
        for row in self.table.rows:
            for cell in row.cells:
                yield cell

    def column_cells(self, col_index: int) -> iter:
        """Итератор по ячейкам колонки."""
        for row in self.rows:
            yield row.cells[col_index]

    def row_cells(self, row_index: int) -> iter:
        """Итератор по ячейкам строки"""
        row = self.table.rows[row_index]
        yield from row.cells

    def iter_cols_without_header(self, col_index: int) -> list:
        """
        Возвращает список с содержимым ячеек указанной колонки, без заголовка
        """
        cells = enumerate(self.column_cells(col_index))
        return [cell for index, cell in cells if index != 0]

    @staticmethod
    def set_cell_value(cell: _Cell, value: Any) -> None:
        """Сеттер, устанавливает значение в ячейке."""
        cell.text = str(value)
        styler.apply_cell_style(cell)

    def set_row_height(self, row_index: int, cm: int | float) -> None:
        """Сеттер, устанавливает высоту ячеек в строке (в сантиметрах)."""
        row = self.get_row(row_index)
        row.height = Cm(cm)

    @staticmethod
    def set_cell_alignment(cell: _Cell, align: str = None) -> None:
        """Сеттер, устанавливает выравнивание текста в ячейке."""
        if "left" in align.lower():
            align_code = styler.align.LEFT
        elif "center" in align.lower():
            align_code = styler.align.CENTER
        elif "right" in align.lower():
            align_code = styler.align.RIGHT
        else:
            align_code = styler.align.JUSTIFY
        cell.paragraphs[0].alignment = align_code

    @staticmethod
    def set_cell_vertical_alignment(cell: _Cell, align: str) -> None:
        """Сеттер, устанавливает вертикальное выравнивание ячейки."""
        if "top" in align.lower():
            align_code = WD_CELL_VERTICAL_ALIGNMENT.TOP
        elif "bottom" in align.lower():
            align_code = WD_CELL_VERTICAL_ALIGNMENT.BOTTOM
        elif "both" in align.lower():
            align_code = WD_CELL_VERTICAL_ALIGNMENT.BOTH
        else:
            align_code = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cell.vertical_alignment = align_code

    @staticmethod
    def set_cell_bold(cell: _Cell) -> None:
        """Превращает текст в ячейке в жирный."""
        cell.paragraphs[0].runs[0].font.bold = True

    @staticmethod
    def set_cell_italic(cell: _Cell) -> None:
        """Превращает текст в ячейке в курсив."""
        cell.paragraphs[0].runs[0].font.italic = True

    def page_break(self):
        """Вставка разрыва страницы."""
        self.doc.add_page_break()

    def section_break(self):
        """Вставка разрыва раздела (на следующей странице)."""
        self.doc.add_section(WD_SECTION.NEW_PAGE)

    def section_break_current(self):
        """Вставка разрыва раздела (на текущей странице)."""
        self.doc.add_section(WD_SECTION.CONTINUOUS)

    def new_column_break(self):
        self.doc.add_section(WD_SECTION.NEW_COLUMN)

    def set_column_width(self, col_index, cm: int | float):
        """
        Сеттер, устанавливает ширину колонки
        Передаем параметр cm в сантиметрапх
        """
        for col in self.column_cells(col_index):
            col.width = Cm(cm)

    def get_cell(self, row_index: int, column_index: int) -> _Cell:
        """Геттер, возвращает ячейку по ее индексам."""
        return self.table.cell(row_index, column_index)

    def set_global_font_style(self, name: str, size: int):
        """Сеттер, устанавливает глобальный шрифт и его размер."""
        styler.reset_global_font_style(self.doc, name, size)

    @staticmethod
    def _set_repeat_table_header(row) -> None:
        """Сеттер, устанавливает повторяющиеся строки таблицы."""
        tr = row._tr
        trPr = tr.get_or_add_trPr()
        table_header = OxmlElement("w:tblHeader")
        table_header.set(qn("w:val"), "true")
        trPr.append(table_header)

    def _build_table(
        self, rows: int = None, cols: int = None, table_style: str = None
    ) -> Table:
        """Метод создает таблицу и добавляет в документ."""
        return self.doc.add_table(
            rows=rows if rows else len(self.data),
            cols=cols if cols else len(self.data[0]),
            style=table_style,
        )

    def _fill_table(self) -> None:
        """Метод заполняет таблицу данными."""
        for cell_data, cell in zip(self._data, self.cells):
            self.set_cell_value(cell, cell_data)

    def repeat_header(self, row: int = 0, rows: list[int] = None):
        """Метод устанавливает повторяющиеся заголовки или один заголовок."""
        if rows:
            for row in rows:
                self._set_repeat_table_header(self.get_row(row))
        else:
            self._set_repeat_table_header(self.get_row(row))

    def merge_cells(
        self, start_cell: dict[str:int, str:int], end_cell: dict[str:int, str:int]
    ) -> None:
        """
        Метод сливает ячейки.
        Значения индексов ячеек нужно передавать словарем вида
        {'col': 0, 'row': 0}
        """
        Cell = namedtuple(
            "C1",
            "col, row",
        )
        start_merge_cell = Cell(col=start_cell["col"], row=start_cell["row"])
        end_merge_cell = Cell(col=end_cell["col"], row=end_cell["row"])
        self._merge(
            CellSelection(
                c1=start_merge_cell,
                c2=end_merge_cell,
                max_row=self.last_row,
                max_col=self.last_col,
            )
        )

    def round_col(self, index: int, round_to: int, skip_rows: list[int] = None) -> None:
        """Метод округляет числа в колонке."""
        for idx, row in enumerate(self.rows):
            if skip_rows and idx in skip_rows:
                continue
            cell = row.cells[index]
            try:
                cell_text = cell.text
                value = Decimal(cell_text)
            except decimal.InvalidOperation:
                continue
            self.set_cell_value(cell, round(value, round_to))

    def add_total(self, cols_index: list, totals: list) -> None:
        """
        Метод добавляет строку итого, для численной колонки.
        Суммируются значения из списка 'totals' для колонок из списка cols_index.
        """
        new = self.table.add_row()
        new.cells[0].text = "Итого"
        for col, total in zip(cols_index, totals):
            self.set_cell_value(self.get_cell(self.last_row, col), total)

    def replace_char_in_decimals(
        self, index_cols: list[int], old: str, new: str
    ) -> None:
        """Метод заменяет символ в колонках на новый."""
        for index in index_cols:
            for cell in self.column_cells(index):
                try:
                    Decimal(cell.text)
                    self.set_cell_value(cell, cell.text.replace(old, new))
                except decimal.InvalidOperation:
                    continue

    def _merge(self, cells: CellSelection) -> None:
        """Метод сливает колонки. Внутренняя функция."""
        start_cell = self.get_cell(row_index=cells.c1.row, column_index=cells.c1.col)
        stop_cell = self.get_cell(row_index=cells.c2.row, column_index=cells.c2.col)

        cell_value = start_cell.text

        merged_cell = start_cell.merge(stop_cell)

        for cell_paragraph in self.iter_cell_text(merged_cell):
            paragraph_element = cell_paragraph._element
            paragraph_element.getparent().remove(paragraph_element)

        self.set_cell_value(merged_cell, cell_value)

    def set_page_size_and_orient(
        self,
        width: int = 210,
        height: int = 297,
        orientation: str = "portrait",
        table_index: int = None,
    ):
        """
        Сеттер, устаналивает размеры страницы и его ориентацию.
        Параметры width и height задаются в миллиметрах.
        """
        section = self.get_section(-1)
        section.page_height = Mm(height)
        section.page_width = Mm(width)
        if orientation == "portrait":
            section.page_height = Mm(height)
            section.page_width = Mm(width)
        elif orientation == "landscape":
            section.page_height = Mm(width)
            section.page_width = Mm(height)
        if table_index:
            table = self.get_table(table_index)
            page_width = section.page_width - section.left_margin - section.right_margin
            column_width = page_width / len(table.columns)
            for cell in self.cells:
                cell.width = column_width

    def get_last_section(self):
        """Геттер, возвращает последний раздел."""
        return self.get_section(-1)

    def make_text_vertical_in_cell(self, cell: _Cell, position: str = None) -> None:
        """
        Метод переворачивает тектс в ячейке
        с горизонтального на вертикальное положение
        """
        position_map = {"bottom_to_top": "btLr", "top_to_bottom": "lrTb"}
        position_value = position_map.get(position, "btLr")
        tc = cell._element
        text_direction = OxmlElement("w:textDirection")
        text_direction.set(qn("w:val"), position_value)
        tcPr = tc.get_or_add_tcPr()
        tcPr.append(text_direction)

    def delete_table_border(self, row_index: int, position: str) -> None:
        """
        Метод удаляет границу таблицы.
        Пока умеет только верхнюю (top) и нижнюю (bottom).
        """
        xml_str = {
            "bottom": r"<w:tcBorders xmlns:w="
            r'"http://schemas.openxmlformats.org'
            r'/wordprocessingml/2006/main">'
            r'<w:top w:val="single" w:sz="4" w:space="0" />'
            r'<w:left w:val="single" w:sz="4" w:space="0" />'
            r'<w:bottom w:val="nil" />'
            r'<w:right w:val="single" w:sz="4" w:space="0" />'
            r"</w:tcBorders>",
            "top": r"<w:tcBorders xmlns:w="
            r'"http://schemas.openxmlformats.org'
            r'/wordprocessingml/2006/main">'
            r'<w:top w:val="nil" />'
            r'<w:left w:val="single" w:sz="4" w:space="0" />'
            r'<w:bottom w:val="single" w:sz="4" w:space="0" />'
            r'<w:right w:val="single" w:sz="4" w:space="0" />'
            r"</w:tcBorders>",
        }
        for cell in self.get_row(row_index).cells:
            cell._element.get_or_add_tcPr().append(oxml.parse_xml(xml_str[position]))

    def _replace_text(self, searched_string: str, replaced_string: str) -> None:
        """
        Метод заменяет кусок текста из ячейки, с сохранением стилей.
        """
        for row in self.template.rows:
            for cell in row.cells:
                for prg in cell.paragraphs:
                    for run in prg.runs:
                        if searched_string in run.text:
                            run.text = run.text.replace(
                                searched_string, replaced_string
                            )

    def copy_table(
        self,
        source_path: str,
        destination_path: str,
        table_index: int,
        searched: str = None,
        new_text: str = None,
    ) -> None:
        """
        Метод копирует таблицу по индексу table_index
        Из заданного документа.
        При необходимости находит строку и заменяет ее.
        """
        self.source_document = Document(source_path)
        self.template = self.source_document.tables[table_index]
        if searched:
            self._replace_text(searched_string=searched, replaced_string=new_text)
        tbl = self.template._tbl
        new_tbl = deepcopy(tbl)
        self.destination_document = self.doc
        self.destination_document.add_paragraph()._p.addnext(new_tbl)
        self.destination_document.save(destination_path)

    def copy_text(
        self,
        source_path: str,
        destination_path: str,
        start_text: str,
        delete_start_text: bool = False,
    ) -> None:
        """
        Метод копирует текст из заданного документа,
        начиная с определенного словосочения и заканчивая
        при нахождении другого словосочения (если указано).
        """
        self.source_document = Document(source_path)
        self.destination_document = self.doc
        copying_flag = False

        for paragraph in self.source_document.paragraphs:
            if paragraph.text.startswith(start_text):
                copying_flag = True
            if copying_flag:
                if paragraph.style.name == "Heading 1" or paragraph.text.strip() == "":  # noqa: E501
                    break

                new_paragraph = self.destination_document.add_paragraph()
                new_paragraph.style = paragraph.style
                self.copy_paragraph_attributes(paragraph, new_paragraph)

                for run in paragraph.runs:
                    if delete_start_text and start_text in paragraph.text:
                        continue

                    new_run = new_paragraph.add_run(run.text)
                    new_run.bold = run.bold
                    new_run.italic = run.italic
                    new_run.underline = run.underline
                    new_run.font.color.rgb = run.font.color.rgb
                    new_run.font.size = run.font.size
                    new_run.font.name = run.font.name
        self.destination_document.save(destination_path)

    def copy_paragraph_attributes(
        self, source_paragraph: Paragraph, destination_paragraph: Paragraph
    ) -> None:  # какое-то хреновое описание
        """
        Метод копирует все параметры текста (отступы, выравнивание и т.д.),
        который нужно вставить в документ.
        """
        if source_paragraph.alignment is not None:
            destination_paragraph.alignment = source_paragraph.alignment
        link = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        num_pr = source_paragraph._element.find(f".//{{{link}}}numPr")
        if num_pr is not None:
            destination_paragraph._element.get_or_add_pPr().append(num_pr)
        if source_paragraph.alignment is not None:
            destination_paragraph.alignment = source_paragraph.alignment
        if source_paragraph.paragraph_format.left_indent is not None:
            left_indent = source_paragraph.paragraph_format.left_indent
            destination_paragraph.paragraph_format.left_indent = left_indent
        if source_paragraph.paragraph_format.right_indent is not None:
            right_indent = source_paragraph.paragraph_format.right_indent
            destination_paragraph.paragraph_format.right_indent = right_indent
        if source_paragraph.paragraph_format.space_before is not None:
            space_before = source_paragraph.paragraph_format.space_before
            destination_paragraph.paragraph_format.space_before = space_before
        if source_paragraph.paragraph_format.space_after is not None:
            space_after = source_paragraph.paragraph_format.space_after
            destination_paragraph.paragraph_format.space_after = space_after

    def merge_depending_on_2columns(
        self, base_col_idx_1: int, merge_col_idx: int, base_col_idx_2: int = None
    ) -> list[list]:
        """
        Метод объединяет ячейки колонки в рамках заданой базовой колонки.
        """
        pred_cell = None
        merge_row_indexes = list()
        for idx, row in enumerate(self.table.rows):
            if idx > 1:
                if base_col_idx_2:
                    curr_cell = f"{row.cells[base_col_idx_1].text}_{row.cells[base_col_idx_2].text}_{row.cells[merge_col_idx].text}"  # noqa: E501
                else:
                    curr_cell = f"{row.cells[base_col_idx_1].text}_{row.cells[merge_col_idx].text}"  # noqa: E501
                if pred_cell and curr_cell == pred_cell:
                    if merge_row_indexes and idx - 1 in merge_row_indexes[-1]:
                        merge_row_indexes[-1].pop()
                        merge_row_indexes[-1].append(idx)
                    else:
                        merge_row_indexes.append([idx - 1, idx])
                pred_cell = curr_cell
        return merge_row_indexes

    def update_section_columns(self, column_nums: int) -> None:
        """
        Метод делит на колонки страницы в последнем разделе.
        """
        link = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        section = self.get_section(-1)
        section_xml = section._sectPr
        cols_elem = section_xml.find(f".//{{{link}}}cols")
        if cols_elem is not None:
            cols_elem.set(f"{{{link}}}num", str(column_nums))
            if column_nums > 1:
                for _ in range(column_nums):
                    col = etree.Element(f"{{{link}}}col")
                    cols_elem.append(col)
            section_xml.append(cols_elem)
