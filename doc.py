import os
import tempfile
from io import BytesIO

# from docx.api import Document
from docx.document import Document
from docx.section import Section
from docx.shared import Cm
from docx.table import Table
from docx2pdf import convert



from docx.package import Package
from docx.opc.package import OpcPackage
from docx.parts.document import DocumentPart
from typing import IO, TYPE_CHECKING, cast
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.document import Document as DocumentObject
from docx.oxml.xmlchemy import BaseOxmlElement, ZeroOrMore, ZeroOrOne


class DOC(Document):
    _left_margin: int = 3
    _right_margin: int = 1.5

    def __init__(self, template: str = None):
        # Инициализация Package
        # Загрузка документа
        document_part = Package.open(template).main_document_part
        document = document_part.document
        element = getattr(document, '_element', None)
        Document.__init__(self, element, document_part)



    # def __init__(self, template: str = None):
    #
    #     document_part = cast("DocumentPart", Package.open(template).main_document_part)
    #     document = document_part.document
    #     self._part = document.__dict__.get('_part')
    #     self._element = document.__dict__.get('_element')
    #     super(Package, self).__init__()
    #     super(DocumentObject, self).__init__()


        # self.doc = DOC.create_document(template)

        # self.setup_margins()

    def _default_docx_path(self):
        """Return the path to the built-in default .docx package."""
        _thisdir = os.path.split(__file__)[0]
        return os.path.join(_thisdir, "templates", "default.docx")


    def get_table(self, table_index: int) -> Table:
        """Возвращает таблицу, по индексу."""
        return self.doc.tables[table_index]

    def get_section(self, section_index: int) -> Section:
        """
        Возвращает секцию, по индексу.
        Пригождается при изменении ориентации страницы.
        """
        return self.doc.sections[section_index]

    @property
    def iter_sections(self) -> iter:
        """Итератор по секциям."""
        yield from self.doc.sections

    def setup_margins(self) -> None:
        """Устанавливает отступы документа слева и справа."""
        for section in self.iter_sections:
            section.left_margin = Cm(self._left_margin)
            section.right_margin = Cm(self._right_margin)

    @staticmethod
    def create_document(template_path: str) -> Document:  ############
        """
        Создает документ из шаблона. В шаблоне нужно определить стили.
        """
        return Document(template_path)

    def delete_empty_page_if_exist(self) -> None:
        """
        Вызывается перед созданием байтового представления документа
        Метод проверяет наличие пустой страницы в конце,
        которая могла появиться из-за разрывов страниц, и удаляет ее.
        """
        paragraphs = self.doc.paragraphs
        if paragraphs:
            last_paragraph = paragraphs[-1]
            if not last_paragraph.text.strip():
                p = last_paragraph._element
                p.getparent().remove(p)

    @property
    def doc_bytes(self) -> bytes:
        """
        Метод возвращает байтовое представление документа,
        которое далее используется либо для передачи документа
        по HTTP, либо для записи в файл.
        """
        buffer = BytesIO()
        self.save(buffer)
        content = buffer.getvalue()
        return content

    def save_doc(self, output_docx_path: str) -> None:
        """
        Метод сохраняет документ в директорию проекта.
        Использовать при разработке, для визуального отображения результата.
        """
        if output_docx_path.split(".")[-1] == "docx":
            with open(output_docx_path, "wb") as file:
                file.write(self.doc_bytes)
        else:
            raise ValueError("'output_docx_path' is not .docx file")

    @staticmethod
    def convert_docx_to_pdf(dirs: str):  ###
        """
        использует библиотеку docx2pdf
        :dirs - путь к папке с docx-файлами, которые конвертируются в pdf
        """
        files_for_convert = os.listdir(dir)
        pass

    def save_pdf(self, output_pdf_path: str) -> None:
        """
        Конвертирует байты DOCX в PDF-файл.

        :param output_pdf_path: Путь, куда сохранить PDF
        (например, "output.pdf").
        """
        if output_pdf_path.split(".")[-1] == "pdf":
            with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp_docx:
                tmp_docx.write(self.doc_bytes)
                tmp_docx_path = tmp_docx.name

            convert(tmp_docx_path, output_pdf_path)

            os.unlink(tmp_docx_path)
        else:
            raise ValueError("'output_pdf_path' is not .pdf file")

    @staticmethod
    def load_docx_as_bytes() -> bytes:
        doc = Document("error_doc.docx")
        buffer = BytesIO()
        doc.save(buffer)
        doc_bytes = buffer.getvalue()
        return doc_bytes

    def add_table(self, rows: int, cols: int, style: str | _TableStyle | None = None):
        """Add a table having row and column counts of `rows` and `cols` respectively.

        `style` may be a table style object or a table style name. If `style` is |None|,
        the table inherits the default table style of the document.
        """
        table = self._body.add_table(rows, cols, self._block_width)
        table.style = style
        return table

class Tabel:

    def __init__(self):
        super().__init__()


d = DOC(r'D:\projects\docx_gen\venv\Lib\site-packages\docx\templates\default.docx')
# print(d._element)

d.add_table(5, 6)
print(d.tables)
d.save_doc(r'C:\Users\Krasnopoliskij.IA\Downloads\Новая папка (36)\hui.docx')
